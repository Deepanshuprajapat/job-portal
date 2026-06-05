import os
import sqlite3
import docx
from app import app, init_db, calculate_match_score, extract_text_from_docx, get_db

def create_test_docx(filepath):
    """Helper to generate a dummy DOCX resume for testing."""
    doc = docx.Document()
    doc.add_heading('Deepanshuprajapat - Software Engineer', 0)
    doc.add_paragraph('Skills: Python, Flask, SQLite, HTML, CSS, JavaScript')
    doc.add_paragraph('Experience: Developed full-stack web applications using Flask backend and SQLite database.')
    doc.add_paragraph('Education: Bachelor of Computer Science.')
    doc.save(filepath)
    print(f"[TEST] Created dummy docx resume at {filepath}")

def run_tests():
    print("[TEST] Initializing database...")
    db_path = app.config.get('DATABASE') or os.path.join(app.root_path, 'database.db')
    if os.path.exists(db_path):
        os.remove(db_path)
        print("[TEST] Removed old database.")

    # Initialize DB
    init_db()

    with app.app_context():
        db = get_db()
        
        # Test 1: Register users
        print("[TEST] Testing User Registration...")
        db.execute(
            "INSERT INTO users (username, email, password_hash, role) VALUES (?, ?, ?, ?)",
            ("test_recruiter", "recruiter@test.com", "hash123", "recruiter")
        )
        db.execute(
            "INSERT INTO users (username, email, password_hash, role) VALUES (?, ?, ?, ?)",
            ("test_candidate", "candidate@test.com", "hash456", "candidate")
        )
        db.commit()
        
        recruiter = db.execute("SELECT * FROM users WHERE username = 'test_recruiter'").fetchone()
        candidate = db.execute("SELECT * FROM users WHERE username = 'test_candidate'").fetchone()
        
        assert recruiter is not None, "Recruiter insertion failed"
        assert candidate is not None, "Candidate insertion failed"
        print(f"[TEST] Users registered successfully. Recruiter ID: {recruiter['id']}, Candidate ID: {candidate['id']}")
        
        # Test 2: Post a job
        print("[TEST] Testing Job Posting...")
        db.execute('''
            INSERT INTO jobs (recruiter_id, title, company, location, description, requirements)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (
            recruiter['id'], 
            "Python Backend Developer", 
            "AITech Labs", 
            "Remote", 
            "Looking for a backend developer who is experienced in building REST APIs and web servers using Flask and Python. Knowledge of SQLite and HTML/CSS is a plus.",
            "Python, Flask, SQLite"
        ))
        db.commit()
        
        job = db.execute("SELECT * FROM jobs WHERE recruiter_id = ?", (recruiter['id'],)).fetchone()
        assert job is not None, "Job posting failed"
        print(f"[TEST] Job posted successfully: '{job['title']}' at '{job['company']}'")
        
        # Test 3: Parse docx and calculate similarity score
        print("[TEST] Testing DOCX parser and Match Score...")
        docx_path = "test_resume.docx"
        create_test_docx(docx_path)
        
        extracted_text = extract_text_from_docx(docx_path)
        assert "Python" in extracted_text, "DOCX text extraction failed"
        print("[TEST] DOCX text extracted successfully.")
        
        # Calculate score
        score = calculate_match_score(extracted_text, job['requirements'], job['description'])
        print(f"[TEST] Calculated Resume Match Score: {score}%")
        assert score > 0.0, "Similarity score calculation returned 0 or error"
        
        # Clean up docx file
        if os.path.exists(docx_path):
            os.remove(docx_path)
            print("[TEST] Cleaned up dummy docx resume.")

        # Test 4: Submit Application
        print("[TEST] Testing Application submission...")
        db.execute('''
            INSERT INTO applications (job_id, candidate_id, resume_filename, match_score, parsed_text)
            VALUES (?, ?, ?, ?, ?)
        ''', (job['id'], candidate['id'], "resume_test.docx", score, extracted_text))
        db.commit()
        
        application = db.execute("SELECT * FROM applications WHERE candidate_id = ?", (candidate['id'],)).fetchone()
        assert application is not None, "Application submission failed"
        assert application['match_score'] == score, "Application match score mismatch"
        print(f"[TEST] Application submitted successfully with match score {application['match_score']}%")
        
        # Test 5: Update Application Status
        print("[TEST] Testing recruiter application status update...")
        db.execute(
            "UPDATE applications SET status = ? WHERE id = ?",
            ("Shortlisted", application['id'])
        )
        db.commit()
        
        updated_app = db.execute("SELECT * FROM applications WHERE id = ?", (application['id'],)).fetchone()
        assert updated_app['status'] == "Shortlisted", "Status update failed"
        print(f"[TEST] Application status updated successfully to: {updated_app['status']}")
        
    print("\n[SUCCESS] All verification tests passed successfully!")

if __name__ == "__main__":
    run_tests()
